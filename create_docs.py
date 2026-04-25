from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x34, 0x4A, 0x6B)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x6B, 0x8D)

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F4F8')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

def tip(text, label="TIP"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(6)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  'E8F5E9')
    p._p.get_or_add_pPr().append(shading)
    r1 = p.add_run(f"✅ {label}: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    p.add_run(text)

def warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(6)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  'FFF8E1')
    p._p.get_or_add_pPr().append(shading)
    r1 = p.add_run("⚠️  NOTE: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
    p.add_run(text)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'),   'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'),  '1A73E8')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'),   'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'),  fill)
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("Google Ads Automation System")
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Complete User Guide & Reference Manual")
r2.font.size  = Pt(16)
r2.font.color.rgb = RGBColor(0x55, 0x6B, 0x8D)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = desc.add_run(
    "Automate bids · Create campaigns · Manage budgets · Get reports\n"
    "Built with Python · FastAPI · Google Ads API · APScheduler"
)
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = date_p.add_run("April 2026  |  sarthak.growleads@gmail.com")
r4.font.size = Pt(10)
r4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_items = [
    ("1",   "What Is This System?"),
    ("2",   "How to Start the System"),
    ("3",   "Dashboard Overview"),
    ("4",   "Automation Rules — The Core Engine"),
    ("5",   "Bid Management"),
    ("6",   "Campaign Creation (CSV & Google Sheets)"),
    ("7",   "Budget Management"),
    ("8",   "Reports & Alerts"),
    ("9",   "Jobs — Running & Monitoring Automation"),
    ("10",  "Email & Slack Notifications Setup"),
    ("11",  ".env Configuration Reference"),
    ("12",  "CSV Campaign Template"),
    ("13",  "Automation Rule Examples"),
    ("14",  "Troubleshooting"),
    ("15",  "Quick Reference Card"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"  {num}.  ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    p.add_run(title)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  1. WHAT IS THIS SYSTEM
# ══════════════════════════════════════════════════════════════
h1("1. What Is This System?")
body(
    "The Google Ads Automation System is a Python-based platform that connects directly "
    "to your Google Ads account and automates four key tasks:"
)
bullet("Bid Management — automatically raises or lowers keyword bids based on CPA, CTR, device, or time of day.", "🎯")
bullet("Campaign Creation — create dozens of campaigns at once by uploading a CSV file or linking a Google Sheet.", "📋")
bullet("Budget Management — monitors monthly spend, pauses campaigns that hit their cap, and redistributes budget.", "💰")
bullet("Reporting & Alerts — sends daily and weekly performance reports to your email and Slack.", "📊")

doc.add_paragraph()
body("The system runs as a local web application on your computer. Once started, it:")
bullet("Runs automation jobs automatically in the background (every 1–2 hours)")
bullet("Provides a web dashboard you open in your browser at http://localhost:8000")
bullet("Sends you alerts via email and Slack when something important happens")
bullet("Keeps a full audit log of every change made to your campaigns")

tip("You do NOT need to write any code to use this system. Everything is controlled through the web dashboard.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  2. HOW TO START
# ══════════════════════════════════════════════════════════════
h1("2. How to Start the System")

h2("Every Time You Want to Use It")
body("Open a Command Prompt or PowerShell window, then run these two commands:")
code("cd \"C:\\Users\\HP\\Desktop\\Google Ads automation\"")
code(".venv\\Scripts\\python.exe run.py")
body("Then open your browser and go to:")
code("http://localhost:8000")
tip("Bookmark http://localhost:8000 in your browser for quick access.")

h2("How to Stop the System")
body("Go back to the terminal window and press Ctrl + C. This stops the server and all automation jobs.")

h2("How to Start Automatically on Windows Startup")
numbered("Press Win + R, type shell:startup, press Enter")
numbered("Create a new .bat file with this content:")
code(
    "@echo off\n"
    "cd \"C:\\Users\\HP\\Desktop\\Google Ads automation\"\n"
    ".venv\\Scripts\\python.exe run.py"
)
numbered("Save it as start_ads_automation.bat in the Startup folder")
body("Now the system starts automatically every time Windows boots.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  3. DASHBOARD OVERVIEW
# ══════════════════════════════════════════════════════════════
h1("3. Dashboard Overview")
body("The dashboard has 6 pages accessible from the top navigation bar:")

add_table(
    ["Page", "URL", "What You Can Do"],
    [
        ["Dashboard",  "http://localhost:8000/",               "See today's KPIs, 7-day trends, recent jobs"],
        ["Campaigns",  "http://localhost:8000/campaigns",      "View all campaigns, upload CSV, import Google Sheets"],
        ["Budgets",    "http://localhost:8000/budgets",        "Set monthly caps, view spend, pause/unpause campaigns"],
        ["Rules",      "http://localhost:8000/rules",          "Create, edit, enable/disable automation rules"],
        ["Reports",    "http://localhost:8000/reports/daily",  "View daily and weekly performance charts"],
        ["Jobs",       "http://localhost:8000/jobs",           "Trigger jobs manually, view audit log"],
    ],
    col_widths=[1.2, 2.2, 2.8]
)

h2("Dashboard KPI Cards")
body("The main dashboard shows 6 KPI cards for today's performance:")
add_table(
    ["Card", "What It Means"],
    [
        ["Impressions",  "How many times your ads were shown today"],
        ["Clicks",       "How many times people clicked your ads today"],
        ["Spend ($)",    "Total money spent on ads today"],
        ["Conversions",  "Number of conversions (purchases, signups, etc.) today"],
        ["CTR (%)",      "Click-through rate = Clicks ÷ Impressions × 100"],
        ["CPA ($)",      "Cost per conversion = Spend ÷ Conversions"],
    ],
    col_widths=[1.5, 4.5]
)
tip("The 7-day trend charts below the KPI cards show whether your performance is improving or declining.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  4. AUTOMATION RULES
# ══════════════════════════════════════════════════════════════
h1("4. Automation Rules — The Core Engine")
body(
    "Rules are the heart of the automation system. Each rule watches your campaign performance "
    "and takes automatic action when specific conditions are met. You create rules once, "
    "and the system evaluates them every 15 minutes to 2 hours automatically."
)

h2("Rule Structure")
body("Every rule has these parts:")
add_table(
    ["Part", "What It Does", "Example"],
    [
        ["Name",        "Label for the rule (you choose)",              "Pause high-CPA keywords"],
        ["Rule Type",   "What kind of automation it does",              "cpa_bid"],
        ["Scope",       "Account-wide, or one specific campaign/keyword", "account"],
        ["Conditions",  "When should the rule fire? (JSON)",            "CPA > $20 AND clicks >= 30"],
        ["Action",      "What should happen when conditions are met?",  "Pause the keyword"],
        ["Priority",    "Lower number = runs first (1–100)",            "1"],
        ["Active",      "Toggle ON/OFF without deleting the rule",      "ON"],
    ],
    col_widths=[1.2, 2.5, 2.5]
)

h2("Rule Types")
add_table(
    ["Rule Type",                "What It Controls",                    "Runs Every"],
    [
        ["cpa_bid",               "Bid up/down based on cost-per-conversion", "1 hour"],
        ["tod_bid",               "Bid adjustments by time of day",           "15 minutes"],
        ["device_bid",            "Bid adjustments by device (mobile/desktop)", "1 hour"],
        ["keyword_performance_bid","Bid or pause based on keyword metrics",   "1 hour"],
        ["budget_cap",            "Enforce monthly spend cap per campaign",   "2 hours"],
        ["budget_pause",          "Pause campaign when budget rule triggers", "2 hours"],
        ["budget_redistribute",   "Move budget from weak to strong campaigns","2 hours"],
    ],
    col_widths=[2.0, 2.5, 1.5]
)

h2("Condition Metrics")
body("You can build conditions using any of these metrics:")
add_table(
    ["Metric",          "Meaning",                            "Example Value"],
    [
        ["cpa",           "Cost per conversion in dollars",     "20.0"],
        ["ctr",           "Click-through rate (0.0 to 1.0)",   "0.02  (= 2%)"],
        ["clicks",        "Number of clicks",                   "50"],
        ["cost",          "Total spend in dollars",             "100.0"],
        ["conversions",   "Number of conversions",              "5"],
        ["roas",          "Return on ad spend",                 "3.0  (= 300%)"],
        ["cpc",           "Average cost per click",             "1.50"],
        ["hour_of_day",   "Current hour (0–23 UTC)",            "9  (= 9am UTC)"],
        ["day_of_week",   "Day 0=Monday … 6=Sunday",            "5  (= Saturday)"],
        ["device",        "Device type",                        "MOBILE / DESKTOP / TABLET"],
    ],
    col_widths=[1.5, 2.8, 1.8]
)

h2("Condition Operators")
add_table(
    ["Operator", "Meaning",         "Example"],
    [
        ["gt",      "Greater than",   "cpa > 20"],
        ["lt",      "Less than",      "ctr < 0.01"],
        ["gte",     "Greater or equal","clicks >= 50"],
        ["lte",     "Less or equal",  "cost <= 500"],
        ["eq",      "Equals",         "device = MOBILE"],
        ["between", "Within a range", "hour_of_day between [9, 18]"],
    ],
    col_widths=[1.2, 1.8, 3.0]
)

h2("Action Types")
add_table(
    ["Action Type",          "What Happens",                             "Required Fields"],
    [
        ["adjust_bid_percent",  "Raise or lower bid by a % (e.g. -10%)",  "value, min_bid_micros, max_bid_micros"],
        ["set_bid_micros",      "Set bid to exact amount in micros",       "value"],
        ["set_bid_adjustment",  "Set device/time bid multiplier",          "value (e.g. 1.2 = +20%)"],
        ["pause_entity",        "Pause the keyword or campaign",           "(none)"],
        ["enable_entity",       "Re-enable a paused keyword/campaign",     "(none)"],
        ["adjust_budget_percent","Change daily budget by a %",            "value"],
        ["set_budget_micros",   "Set daily budget to exact micros",        "value"],
    ],
    col_widths=[1.8, 2.5, 2.0]
)
warn("Bid amounts use MICROS — 1 dollar = 1,000,000 micros. So $0.50 = 500,000 micros.")

h2("How to Create a Rule")
numbered("Go to http://localhost:8000/rules")
numbered("Click the blue '+ New Rule' button")
numbered("Fill in: Name, Rule Type, Scope, Priority")
numbered("In the Conditions box — type your conditions as JSON (see examples below)")
numbered("In the Action box — type your action as JSON")
numbered("Make sure 'Active' is checked")
numbered("Click 'Save Rule'")
tip("Use the Example Rules shown at the bottom of the form page as a starting template.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  5. BID MANAGEMENT
# ══════════════════════════════════════════════════════════════
h1("5. Bid Management")
body(
    "Bid management automatically adjusts how much you pay per click for keywords, "
    "based on their performance. The system runs every hour and checks all your rules."
)

h2("How CPA-Based Bidding Works")
numbered("The system pulls the last 7 days of keyword performance from Google Ads")
numbered("For each keyword, it calculates: CPA = Total Cost ÷ Conversions")
numbered("It checks your rules — if a keyword's CPA is too high, it lowers the bid")
numbered("If the CPA is low (performing well), it can raise the bid to get more traffic")
numbered("Every change is recorded in the Audit Log")

h2("How Time-of-Day Bidding Works")
body(
    "You can increase bids during your best hours and decrease them during slow hours. "
    "This rule runs every 15 minutes and checks the current UTC hour."
)
bullet("Example: Boost bids by +25% during business hours (9am–6pm)")
bullet("Example: Reduce bids by -30% at night (11pm–6am) to save budget")

h2("How Device Bidding Works")
body(
    "If mobile users convert less than desktop users, you can automatically reduce "
    "the bid multiplier for mobile to spend less on low-performing devices."
)

h2("Bid Limits (Important)")
body("Always set min and max bid limits in your rules to prevent extreme bids:")
code('{"type": "adjust_bid_percent", "value": -15, "min_bid_micros": 100000, "max_bid_micros": 5000000}')
body("This means: reduce bid by 15%, but never go below $0.10 (100,000 micros) or above $5.00 (5,000,000 micros)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  6. CAMPAIGN CREATION
# ══════════════════════════════════════════════════════════════
h1("6. Campaign Creation (CSV & Google Sheets)")
body(
    "Instead of creating campaigns one by one in Google Ads, you can create dozens at once "
    "by uploading a spreadsheet. The system reads each row and creates the full campaign "
    "structure: Campaign → Ad Group → Ad → Keywords."
)

h2("Method 1: Upload a CSV File")
numbered("Go to http://localhost:8000/campaigns")
numbered("Click the blue 'Upload CSV' button")
numbered("Select your CSV file and click 'Upload & Create'")
numbered("A batch job starts in the background — campaigns are created automatically")
numbered("Refresh the page to see the status (pending → processing → completed)")

h2("Method 2: Import from Google Sheets")
numbered("Go to http://localhost:8000/campaigns")
numbered("Scroll down to 'Or import from Google Sheets'")
numbered("Paste your Google Sheet URL and click 'Import'")
warn("For Google Sheets import, you need a Service Account JSON file (optional feature). CSV upload works without it.")

h2("CSV File Format")
body("Your CSV file must have these columns (column names must match exactly):")
add_table(
    ["Column Name",        "Required?", "Description",                          "Example"],
    [
        ["campaign_name",    "YES",      "Unique name for the campaign",         "Summer Sale 2026"],
        ["daily_budget_usd", "YES",      "Daily budget in dollars",              "20.00"],
        ["headline_1",       "YES",      "First ad headline (max 30 chars)",     "Buy Shoes Online"],
        ["headline_2",       "YES",      "Second ad headline (max 30 chars)",    "Best Deals Today"],
        ["headline_3",       "YES",      "Third ad headline (max 30 chars)",     "Free Shipping"],
        ["description_1",    "YES",      "First description (max 90 chars)",     "Shop our wide range of shoes"],
        ["description_2",    "YES",      "Second description (max 90 chars)",    "Order today, get free delivery"],
        ["final_url",        "YES",      "Landing page URL",                     "https://example.com/shoes"],
        ["campaign_type",    "No",       "SEARCH or DISPLAY (default: SEARCH)",  "SEARCH"],
        ["bid_strategy",     "No",       "MANUAL_CPC or MAXIMIZE_CONVERSIONS",   "MANUAL_CPC"],
        ["ad_group_name",    "No",       "Ad group name (default: campaign name)","Shoes Ad Group"],
        ["keyword_1",        "No",       "First keyword",                        "buy shoes online"],
        ["keyword_1_match",  "No",       "BROAD, PHRASE, or EXACT",              "BROAD"],
        ["keyword_2",        "No",       "Second keyword",                       "best running shoes"],
        ["keyword_2_match",  "No",       "Match type for keyword 2",             "PHRASE"],
        ["geo_target",       "No",       "Target country (default: US)",         "IN"],
        ["language",         "No",       "Language code (default: en)",          "en"],
    ],
    col_widths=[1.6, 0.8, 2.2, 1.6]
)

tip("A sample CSV file is already in your project folder: sample_campaigns.csv — open it in Excel as a template.")
warn("New campaigns are created in PAUSED status. You must manually enable them in Google Ads after reviewing.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  7. BUDGET MANAGEMENT
# ══════════════════════════════════════════════════════════════
h1("7. Budget Management")
body(
    "The budget management module protects you from overspending. "
    "It runs every 2 hours and checks your campaigns against your defined limits."
)

h2("Monthly Spend Caps")
body(
    "A monthly cap is a maximum amount you allow a campaign to spend in one calendar month. "
    "Once the cap is reached, the campaign is automatically paused."
)

h3("How to Set a Monthly Cap")
numbered("Go to http://localhost:8000/budgets")
numbered("Find the campaign you want to cap")
numbered("In the 'Set Cap' box, enter the monthly limit in dollars (e.g. 500)")
numbered("Click 'Set Cap'")
numbered("A progress bar now shows how much of the cap has been used")
body("The system checks spend every 2 hours. When spent ≥ cap, the campaign is automatically paused and you get a Slack alert.")

h3("How to Re-enable a Paused Campaign")
numbered("Go to http://localhost:8000/budgets")
numbered("Find the campaign showing 'Paused by cap'")
numbered("Click the 'Re-enable' button")
warn("If you re-enable a campaign that hit its monthly cap, it will continue spending. Consider increasing the cap first.")

h2("Budget Rules")
body("You can also create dynamic budget rules (under Rules page):")
bullet("Auto-increase budget for campaigns with ROAS > 5x", "↑")
bullet("Auto-decrease budget for campaigns with zero conversions in 7 days", "↓")
bullet("Redistribute budget from paused campaigns to active ones", "⇄")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  8. REPORTS & ALERTS
# ══════════════════════════════════════════════════════════════
h1("8. Reports & Alerts")

h2("Daily Report")
body("Every day at 6:00 AM UTC, the system automatically:")
bullet("Pulls yesterday's performance from Google Ads")
bullet("Detects anomalies (spend or clicks that deviate >30% from the 7-day average)")
bullet("Sends an HTML email report to your configured email address")
bullet("Sends a Slack message if any anomalies were detected")
body("To view the latest report in the dashboard:")
code("http://localhost:8000/reports/daily")

h2("Weekly Report")
body("Every Monday at 7:00 AM UTC, the system sends a weekly summary with:")
bullet("Top campaigns by cost (last 7 days)")
bullet("Top keywords by cost and performance")
bullet("Conversion trends")
code("http://localhost:8000/reports/weekly")

h2("Anomaly Detection")
body(
    "The system compares today's metrics to the 7-day rolling average. "
    "If any metric deviates by more than 30%, it flags it as an anomaly and alerts you."
)
add_table(
    ["Anomaly Type",   "Example",                                 "Alert Sent"],
    [
        ["Spend spike",  "Today $150, avg $90 — that's +67% spike", "Email + Slack"],
        ["Spend drop",   "Today $20, avg $80 — that's -75% drop",   "Email + Slack"],
        ["Click drop",   "Today 5 clicks, avg 50 — unusual",        "Email + Slack"],
    ],
    col_widths=[1.5, 3.0, 1.5]
)

h2("Manual Report Refresh")
body("You can trigger reports at any time without waiting for the schedule:")
numbered("Go to http://localhost:8000/reports/daily")
numbered("Click the 'Refresh' button at the top right")
numbered("Wait ~30 seconds, then reload the page")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  9. JOBS
# ══════════════════════════════════════════════════════════════
h1("9. Jobs — Running & Monitoring Automation")
body(
    "Jobs are the automation tasks that run in the background. "
    "They run automatically on a schedule, but you can also trigger them manually anytime."
)

h2("Scheduled Jobs")
add_table(
    ["Job Name",         "What It Does",                                  "Auto Schedule"],
    [
        ["bid_management",   "Runs all CPA, device, keyword bid rules",      "Every 1 hour"],
        ["tod_bid_rules",    "Runs time-of-day bid adjustment rules",         "Every 15 minutes"],
        ["budget_management","Checks monthly caps, runs budget rules",        "Every 2 hours"],
        ["daily_reporting",  "Generates daily report, sends email+Slack",     "Daily at 6:00 AM UTC"],
        ["weekly_reporting", "Generates weekly report, sends email",          "Monday 7:00 AM UTC"],
    ],
    col_widths=[1.8, 3.0, 1.6]
)

h2("How to Trigger a Job Manually")
numbered("Go to http://localhost:8000/jobs")
numbered("Click any of the job buttons at the top (e.g. 'bid_management')")
numbered("The job starts immediately — the button disables and the page auto-reloads in 3 seconds")
numbered("Scroll down to see the new run appear in the Job History table")

h2("Reading the Job History Table")
add_table(
    ["Column",         "Meaning"],
    [
        ["Job",          "Which automation job ran"],
        ["Status",       "success (green), failed (red), running (yellow)"],
        ["Triggered By", "scheduler (automatic) or manual_ui (you clicked it)"],
        ["Started",      "When the job started"],
        ["Duration",     "How many seconds it took to complete"],
        ["Rules",        "How many rules were evaluated"],
        ["Actions",      "How many changes were made to Google Ads"],
        ["Errors",       "How many errors occurred (click for details)"],
    ],
    col_widths=[1.5, 4.5]
)

h2("Audit Log")
body(
    "Every single change made to your Google Ads account is recorded in the Audit Log "
    "(bottom of the Jobs page). You can see exactly what was changed, when, and whether it succeeded."
)
tip("If you're unsure why a bid changed, check the Audit Log — it shows the old and new values for every action.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  10. EMAIL & SLACK SETUP
# ══════════════════════════════════════════════════════════════
h1("10. Email & Slack Notifications Setup")

h2("Email Setup (Gmail)")
body("Open your .env file and fill in these values:")
code(
    "EMAIL_BACKEND=smtp\n"
    "SMTP_HOST=smtp.gmail.com\n"
    "SMTP_PORT=587\n"
    "SMTP_USER=your-email@gmail.com\n"
    "SMTP_PASSWORD=your-app-password\n"
    "ALERT_EMAIL_FROM=your-email@gmail.com\n"
    "ALERT_EMAIL_TO=your-email@gmail.com"
)
warn("Do NOT use your regular Gmail password. You need to create an App Password.")
h3("How to get a Gmail App Password:")
numbered("Go to your Google Account → Security")
numbered("Enable 2-Step Verification (required)")
numbered("Search for 'App passwords' in your Google Account settings")
numbered("Create a new app password — select 'Mail' and 'Windows Computer'")
numbered("Copy the 16-character password and paste it as SMTP_PASSWORD in .env")

h2("Slack Setup")
body("To receive Slack alerts, you need a Slack Webhook URL:")
numbered("Go to api.slack.com/apps")
numbered("Click 'Create New App' → 'From Scratch'")
numbered("Name it 'Ads Automation', select your workspace → Create")
numbered("Left menu → 'Incoming Webhooks' → toggle ON")
numbered("Click 'Add New Webhook to Workspace'")
numbered("Choose the channel where you want alerts → Allow")
numbered("Copy the Webhook URL (starts with https://hooks.slack.com/services/...)")
numbered("Paste it in your .env file:")
code("SLACK_WEBHOOK_URL=https://hooks.slack.com/services/T.../B.../...")
body("After saving .env, restart the system (Ctrl+C then python run.py again) for changes to take effect.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  11. .ENV REFERENCE
# ══════════════════════════════════════════════════════════════
h1("11. .env Configuration Reference")
body("Your .env file is at: C:\\Users\\HP\\Desktop\\Google Ads automation\\.env")
body("This file controls all settings. Never share this file — it contains your API credentials.")

add_table(
    ["Variable",                        "Your Value",             "Description"],
    [
        ["GOOGLE_ADS_DEVELOPER_TOKEN",    "SFBXnl9dU…",           "From Google Ads → Admin → API Center"],
        ["GOOGLE_ADS_CLIENT_ID",          "766833233103-…",        "From Google Cloud Console → Credentials"],
        ["GOOGLE_ADS_CLIENT_SECRET",      "GOCSPX-gUx…",          "From Google Cloud Console → Credentials"],
        ["GOOGLE_ADS_REFRESH_TOKEN",      "1//0gCpFpz…",           "Generated by get_refresh_token.py"],
        ["GOOGLE_ADS_LOGIN_CUSTOMER_ID",  "5504503258",            "Your Manager account ID (no dashes)"],
        ["GOOGLE_ADS_TARGET_CUSTOMER_ID", "3673162925",            "The account to automate (no dashes)"],
        ["DATABASE_URL",                  "sqlite:///./ads_automation.db", "Database location (SQLite for local)"],
        ["EMAIL_BACKEND",                 "smtp",                  "smtp or sendgrid"],
        ["SMTP_HOST",                     "smtp.gmail.com",        "Your email server"],
        ["SMTP_PORT",                     "587",                   "587 for Gmail TLS"],
        ["SMTP_USER",                     "you@gmail.com",         "Your email address"],
        ["SMTP_PASSWORD",                 "app-password",          "Gmail App Password (NOT your login password)"],
        ["ALERT_EMAIL_TO",                "you@gmail.com",         "Who receives the reports"],
        ["SLACK_WEBHOOK_URL",             "https://hooks.slack…",  "Slack channel webhook for alerts"],
        ["LOG_LEVEL",                     "INFO",                  "INFO or DEBUG (more detail)"],
    ],
    col_widths=[2.3, 1.5, 2.5]
)

warn("After editing .env, always restart the system: press Ctrl+C in the terminal, then run python run.py again.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  12. CSV TEMPLATE
# ══════════════════════════════════════════════════════════════
h1("12. CSV Campaign Template")
body(
    "Copy this template into Excel or Google Sheets to create your campaigns. "
    "Each row = one campaign. Save as .csv and upload at http://localhost:8000/campaigns"
)
code(
    "campaign_name,campaign_type,daily_budget_usd,bid_strategy,ad_group_name,"
    "keyword_1,keyword_1_match,keyword_2,keyword_2_match,headline_1,headline_2,"
    "headline_3,description_1,description_2,final_url,geo_target,language\n\n"
    "Brand Campaign,SEARCH,50.00,MANUAL_CPC,Brand Terms,"
    "mybrand,EXACT,mybrand reviews,PHRASE,"
    "Official MyBrand Store,Best Prices Guaranteed,Shop Today,"
    "The official store for all MyBrand products.,Get fast delivery on all orders.,"
    "https://example.com,IN,en\n\n"
    "Product Campaign,SEARCH,30.00,MAXIMIZE_CONVERSIONS,Products,"
    "buy product online,BROAD,best product deals,PHRASE,"
    "Buy Products Online,Huge Discounts Today,Free Shipping,"
    "Shop our full range of products now.,Limited time offer - order today!,"
    "https://example.com/products,IN,en"
)
tip("The file 'sample_campaigns.csv' in your project folder is a ready-to-use template.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  13. RULE EXAMPLES
# ══════════════════════════════════════════════════════════════
h1("13. Automation Rule Examples")
body("Copy and paste these directly into the Rules form on the dashboard.")

h2("Example 1: Pause keywords with high CPA")
body("Rule Type: keyword_performance_bid | Scope: account | Priority: 1")
code(
    "Conditions:\n"
    "[\n"
    "  {\"metric\": \"cpa\", \"operator\": \"gt\", \"value\": 25.0},\n"
    "  {\"metric\": \"clicks\", \"operator\": \"gte\", \"value\": 30}\n"
    "]\n\n"
    "Action:\n"
    "{\"type\": \"pause_entity\"}"
)
body("Effect: Any keyword that has had 30+ clicks but a CPA above ₹25 will be paused automatically.")

h2("Example 2: Lower bids on high-CPA keywords (gentler approach)")
body("Rule Type: cpa_bid | Scope: account | Priority: 2")
code(
    "Conditions:\n"
    "[\n"
    "  {\"metric\": \"cpa\", \"operator\": \"gt\", \"value\": 20.0},\n"
    "  {\"metric\": \"clicks\", \"operator\": \"gte\", \"value\": 20}\n"
    "]\n\n"
    "Action:\n"
    "{\n"
    "  \"type\": \"adjust_bid_percent\",\n"
    "  \"value\": -15,\n"
    "  \"min_bid_micros\": 100000,\n"
    "  \"max_bid_micros\": 5000000\n"
    "}"
)
body("Effect: Reduce bid by 15% for any keyword with CPA > ₹20 and at least 20 clicks.")

h2("Example 3: Boost bids during business hours")
body("Rule Type: tod_bid | Scope: account | Priority: 5")
code(
    "Conditions:\n"
    "[\n"
    "  {\"metric\": \"hour_of_day\", \"operator\": \"between\", \"value\": [4, 14]}\n"
    "]\n\n"
    "Action:\n"
    "{\"type\": \"set_bid_adjustment\", \"value\": 1.25}"
)
body("Effect: Apply a +25% bid multiplier from 4am–2pm UTC (9am–7pm India time).")

h2("Example 4: Reduce mobile bids")
body("Rule Type: device_bid | Scope: account | Priority: 3")
code(
    "Conditions:\n"
    "[\n"
    "  {\"metric\": \"device\", \"operator\": \"eq\", \"value\": \"MOBILE\"},\n"
    "  {\"metric\": \"ctr\", \"operator\": \"lt\", \"value\": 0.01}\n"
    "]\n\n"
    "Action:\n"
    "{\"type\": \"set_bid_adjustment\", \"value\": 0.7}"
)
body("Effect: Set mobile bid modifier to 70% (a -30% reduction) for campaigns where mobile CTR is below 1%.")

h2("Example 5: Increase budget for well-performing campaigns")
body("Rule Type: budget_redistribute | Scope: account | Priority: 8")
code(
    "Conditions:\n"
    "[\n"
    "  {\"metric\": \"roas\", \"operator\": \"gt\", \"value\": 4.0},\n"
    "  {\"metric\": \"conversions\", \"operator\": \"gte\", \"value\": 10}\n"
    "]\n\n"
    "Action:\n"
    "{\n"
    "  \"type\": \"adjust_budget_percent\",\n"
    "  \"value\": 20,\n"
    "  \"min_budget_micros\": 500000\n"
    "}"
)
body("Effect: Increase daily budget by 20% for any campaign with ROAS above 4x and at least 10 conversions in 7 days.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  14. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════
h1("14. Troubleshooting")

add_table(
    ["Problem",                               "Cause",                           "Solution"],
    [
        ["Dashboard shows no data",            "API credentials wrong or no campaigns", "Check .env values; verify account has campaigns"],
        ["'Auth failed' error on startup",     "Invalid developer token or refresh token", "Re-run get_refresh_token.py; check token in .env"],
        ["Port 8000 already in use",           "Another app is using port 8000",  "Run: .venv\\Scripts\\python.exe run.py --port 8001"],
        ["Job shows 'failed' in red",          "API error or network issue",      "Click the row, read the error. Check Jobs > Audit Log"],
        ["No email received",                  "Wrong SMTP settings or App Password", "Check SMTP_USER, SMTP_PASSWORD in .env; use Gmail App Password"],
        ["Slack alerts not working",           "Wrong webhook URL",                "Test URL with curl; regenerate webhook in Slack"],
        ["CSV upload fails",                   "Wrong column names or format",     "Check column names match exactly; open sample_campaigns.csv"],
        ["Campaigns created but not visible",  "Created in PAUSED status",         "Go to Google Ads and enable the campaigns manually"],
        ["Rules not firing",                   "Rule is inactive or conditions not met", "Check rule is Active; check metric values vs your thresholds"],
        ["'Module not found' error",           "Running with wrong Python",        "Always use: .venv\\Scripts\\python.exe run.py"],
    ],
    col_widths=[1.8, 1.8, 2.6]
)

h2("How to See Detailed Errors")
numbered("Go to http://localhost:8000/jobs")
numbered("Find the failed job run in the table")
numbered("Look at the Audit Log below — it shows the exact error message per action")
numbered("You can also check the terminal window where you ran python run.py for full stack traces")

h2("Re-generating Your Refresh Token")
body("If your refresh token expires (rare, but can happen):")
numbered("Run:  .venv\\Scripts\\python.exe get_refresh_token.py")
numbered("Complete the browser login again")
numbered("Copy the new token into .env as GOOGLE_ADS_REFRESH_TOKEN")
numbered("Restart the system")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  15. QUICK REFERENCE CARD
# ══════════════════════════════════════════════════════════════
h1("15. Quick Reference Card")

h2("Start / Stop")
add_table(
    ["Action",            "Command / URL"],
    [
        ["Start system",    "cd \"C:\\Users\\HP\\Desktop\\Google Ads automation\"  then  .venv\\Scripts\\python.exe run.py"],
        ["Stop system",     "Press Ctrl + C in the terminal"],
        ["Open dashboard",  "http://localhost:8000"],
    ],
    col_widths=[1.5, 4.8]
)

h2("Key URLs")
add_table(
    ["Page",          "URL"],
    [
        ["Dashboard",   "http://localhost:8000/"],
        ["Campaigns",   "http://localhost:8000/campaigns"],
        ["Budgets",     "http://localhost:8000/budgets"],
        ["Rules",       "http://localhost:8000/rules"],
        ["New Rule",    "http://localhost:8000/rules/new"],
        ["Daily Report","http://localhost:8000/reports/daily"],
        ["Weekly Report","http://localhost:8000/reports/weekly"],
        ["Job History", "http://localhost:8000/jobs"],
    ],
    col_widths=[1.5, 4.8]
)

h2("Important Conversions")
add_table(
    ["Amount",    "In Micros"],
    [
        ["$0.10",   "100,000"],
        ["$0.50",   "500,000"],
        ["$1.00",   "1,000,000"],
        ["$5.00",   "5,000,000"],
        ["$10.00",  "10,000,000"],
        ["$50.00",  "50,000,000"],
    ],
    col_widths=[1.5, 4.8]
)

h2("Your Account IDs")
add_table(
    ["Account",            "ID"],
    [
        ["Manager Account",  "550-450-3258  →  5504503258"],
        ["Ads Account",      "367-316-2925  →  3673162925"],
    ],
    col_widths=[2.0, 4.3]
)

# ── Save ──────────────────────────────────────────────────────
out = r"C:\Users\HP\Desktop\Google Ads automation\Google_Ads_Automation_User_Guide.docx"
doc.save(out)
print(f"Document saved: {out}")
